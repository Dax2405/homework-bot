from docx import Document
from docx.shared import Inches
from news.news_item import NewsItem
import os
from docx2pdf import convert
import platform
import subprocess


def generate_document(national_news_item: NewsItem, international_news_item: NewsItem, editorial_news_item: NewsItem):

    universidad = ""
    materia = ""
    nombre = ""
    curso = ""
    profesor = ""
    path = os.path.join(os.getcwd(), "data/images")

    document = Document()

    document.add_heading(
        f'{universidad}\n{materia}', 1)

    p = document.add_paragraph('')
    p.add_run('Nombre: ').bold = True
    p.add_run(nombre)

    p2 = document.add_paragraph('')
    p2.add_run('Curso: ').bold = True
    p2.add_run(curso)

    p3 = document.add_paragraph('')
    p3.add_run('Profesor: ').bold = True
    p3.add_run(profesor)

    document.add_heading('Noticia Nacional', level=1)
    document.add_picture(
        f'{path}/{national_news_item.title}.jpg', width=Inches(5))

    document.add_paragraph(
        national_news_item.title, style='Heading 2'
    )
    document.add_paragraph(
        national_news_item.opinion
    )

    document.add_heading('Noticia Internacional', level=1)
    international_news_path = f'{path}/{international_news_item.title}.jpg'
    print(international_news_path)
    if os.path.exists(international_news_path):
        document.add_picture(international_news_path, width=Inches(5))
    else:
        print(f"Imagen no encontrada: {international_news_path}")

    document.add_paragraph(
        international_news_item.title, style='Heading 2'
    )
    document.add_paragraph(
        international_news_item.opinion
    )

    document.add_heading('Noticia de Editorial', level=1)
    document.add_picture(
        f'{path}/{editorial_news_item.title}.jpg', width=Inches(5))

    document.add_paragraph(
        editorial_news_item.title, style='Heading 2'
    )
    document.add_paragraph(
        editorial_news_item.opinion
    )

    docx_filename = f'noticias_{nombre}.docx'
    document.save(docx_filename)
    pdf_filename = f'noticias_{nombre}.pdf'
    document.save(pdf_filename)
    if platform.system() == "Windows":
        convert(docx_filename)
    elif platform.system() == "Linux":
        subprocess.run(["libreoffice", "--headless",
                       "--convert-to", "pdf", docx_filename])
    else:
        print("Sistema operativo no soportado para la conversión a PDF")
